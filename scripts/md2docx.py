#!/usr/bin/env python
"""
Markdown to Word (.docx) converter
将 Markdown 报告转换为带格式的 Word 文档

用法:
    python scripts/md2docx.py report/report_cn.md
    python scripts/md2docx.py report/report_en.md
"""

import re
import sys
import os

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """设置单元格背景色。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)


def set_table_style(table):
    """美化表格样式。"""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            # 设置单元格边框
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = OxmlElement('w:tcBorders')
            for edge in ('top', 'left', 'bottom', 'right'):
                element = OxmlElement(f'w:{edge}')
                element.set(qn('w:val'), 'single')
                element.set(qn('w:sz'), '4')
                element.set(qn('w:space'), '0')
                element.set(qn('w:color'), '999999')
                borders.append(element)
            tcPr.append(borders)

            # 表头着色
            if row_idx == 0:
                set_cell_shading(cell, '2F5496')
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True

            # 设置单元格边距
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)


def add_formatted_paragraph(doc, text, style=None, bold=False, font_size=None, color=None, alignment=None):
    """添加格式化段落。"""
    p = doc.add_paragraph()
    if style:
        p.style = doc.styles[style] if style in [s.name for s in doc.styles] else p.style
    run = p.add_run(text)
    if bold:
        run.font.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    return p


def parse_markdown_to_docx(md_path: str, output_path: str):
    """将 Markdown 文件解析并写入 Word 文档。"""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # ── 空行 ──
        if not line:
            i += 1
            continue

        # ── 水平线 ──
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '4')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ── 标题 ──
        header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if header_match:
            level = len(header_match.group(1))
            title_text = header_match.group(2).strip()
            heading = doc.add_heading(title_text, level=level)
            i += 1
            continue

        # ── 代码块 ──
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].rstrip().startswith('```'):
                code_lines.append(lines[i].rstrip())
                i += 1
            i += 1  # 跳过结束的 ```

            # 将代码块添加为灰色背景的段落
            for cl in code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run(cl)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(51, 51, 51)
            continue

        # ── 表格 ──
        if line.startswith('|') and line.endswith('|'):
            # 收集所有表格行
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            # 过滤分隔行 (如 |---|---|)
            data_rows = [tl for tl in table_lines if not re.match(r'^\|[\s\-:|]+\|$', tl)]

            if len(data_rows) < 2:
                continue

            # 解析表头和数据行
            def parse_row(row_str):
                cells = row_str.strip('|').split('|')
                return [c.strip() for c in cells]

            header_cells = parse_row(data_rows[0])
            num_cols = len(header_cells)

            table = doc.add_table(rows=len(data_rows), cols=num_cols)
            table.style = 'Table Grid'

            for row_idx, row_str in enumerate(data_rows):
                cells = parse_row(row_str)
                for col_idx, cell_text in enumerate(cells):
                    if col_idx < num_cols:
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        run = p.add_run(cell_text)
                        if row_idx == 0:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.size = Pt(10)
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            run.font.size = Pt(10)
                        p.paragraph_format.space_before = Pt(1)
                        p.paragraph_format.space_after = Pt(1)

            set_table_style(table)
            doc.add_paragraph()  # 表后空行
            continue

        # ── 有序列表 ──
        list_match = re.match(r'^(\d+)\.\s+(.+)$', line)
        if list_match:
            list_items = []
            while i < len(lines):
                m = re.match(r'^(\d+)\.\s+(.+)$', lines[i].rstrip())
                if not m:
                    break
                list_items.append(m.group(2))
                i += 1

            for item in list_items:
                # 处理粗体标记
                p = doc.add_paragraph(style='List Number')
                _add_runs_with_format(p, item, Pt(11))
            continue

        # ── 无序列表 ──
        bullet_match = re.match(r'^[\-\*]\s+(.+)$', line)
        if bullet_match:
            bullet_items = []
            while i < len(lines):
                m = re.match(r'^[\-\*]\s+(.+)$', lines[i].rstrip())
                if not m:
                    break
                bullet_items.append(m.group(1))
                i += 1

            for item in bullet_items:
                p = doc.add_paragraph(style='List Bullet')
                _add_runs_with_format(p, item, Pt(11))
            continue

        # ── 加粗标题行（以 ** 开头） ──
        bold_header = re.match(r'^\*\*(.+?)\*\*(.*)$', line)
        if bold_header:
            p = doc.add_paragraph()
            run = p.add_run(bold_header.group(1))
            run.font.bold = True
            run.font.size = Pt(12)
            if bold_header.group(2):
                run2 = p.add_run(bold_header.group(2))
                run2.font.size = Pt(11)
            i += 1
            continue

        # ── 普通段落 ──
        p = doc.add_paragraph()
        _add_runs_with_format(p, line, Pt(11))
        i += 1

    # 保存
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)
    print(f"  [OK] {output_path}")


def _add_runs_with_format(paragraph, text: str, font_size: Pt):
    """解析行内格式（粗体、代码）并添加到段落。"""
    # 处理粗体 **text**
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
            run.font.size = font_size
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(font_size.pt - 1)
            run.font.color.rgb = RGBColor(180, 50, 50)
        else:
            # 处理行内公式 $...$
            math_parts = re.split(r'(\$[^$]+\$)', part)
            for mp in math_parts:
                if mp.startswith('$') and mp.endswith('$'):
                    run = paragraph.add_run(mp[1:-1])
                    run.font.italic = True
                    run.font.size = font_size
                else:
                    run = paragraph.add_run(mp)
                    run.font.size = font_size


def main():
    if len(sys.argv) < 2:
        # 默认转换两个文件
        paths = ['report/report_cn.md', 'report/report_en.md']
    else:
        paths = sys.argv[1:]

    print("=" * 60)
    print("  Markdown → Word 转换")
    print("=" * 60)

    for md_path in paths:
        if not os.path.isfile(md_path):
            print(f"  [SKIP] 文件不存在: {md_path}")
            continue

        docx_path = md_path.replace('.md', '.docx')
        print(f"  转换: {md_path} → {docx_path}")
        try:
            parse_markdown_to_docx(md_path, docx_path)
        except Exception as e:
            import traceback
            print(f"    [ERROR] {e}")
            traceback.print_exc()

    print("\n完成!")


if __name__ == '__main__':
    main()
